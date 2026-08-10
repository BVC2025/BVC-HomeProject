"""
Document extractor — turn an uploaded file into plain text + tables.

Every extractor returns the same shape:
  {
    "text": str,                    concatenated plain text
    "page_count": int | None,       for paged formats
    "tables": [                     zero-or-more table blocks
      {
        "sheet":   Optional[str],   for xlsx
        "page":    Optional[int],   for pdf
        "columns": [str],
        "rows":    [[cell, ...]],
      }
    ],
    "has_images": bool,             true if the source contains embedded imagery
    "ocr_applied": bool,            true if we ran tesseract
    "error": Optional[str],         partial-extraction reason
  }

If a heavy optional dep is missing (pdfplumber for PDF tables, pytesseract
for OCR, python-docx for DOCX), the extractor degrades gracefully:
the text portion still comes back, `tables` may be empty, and an
`error` string explains what the extractor couldn't do.
"""

from __future__ import annotations

import io
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional


log = logging.getLogger("hrms_ai.extractor")


# =====================================================================
# Dispatch
# =====================================================================

def sniff_mime(filename: str, given: Optional[str] = None) -> str:
    """Prefer the client-provided MIME type; fall back to extension."""

    if given and "/" in given:
        return given.lower()
    guess, _ = mimetypes.guess_type(filename or "")
    return (guess or "application/octet-stream").lower()


def extract(path: Path, mime: str) -> Dict:
    """Route to the right extractor. `path` is absolute; `mime` is the
    lowercased MIME the caller detected. Always returns a dict of the
    canonical shape — never raises for a corrupt file (the exception
    lands in `error` instead)."""

    m = (mime or "").lower()

    try:
        if "pdf" in m or path.suffix.lower() == ".pdf":
            return _extract_pdf(path)
        if "wordprocessingml" in m or path.suffix.lower() == ".docx":
            return _extract_docx(path)
        if "spreadsheetml" in m or path.suffix.lower() in (".xlsx", ".xlsm"):
            return _extract_xlsx(path)
        if m.startswith("text/") or path.suffix.lower() in (".txt", ".md", ".log", ".csv"):
            if path.suffix.lower() == ".csv":
                return _extract_csv(path)
            return _extract_text(path)
        if m.startswith("image/") or path.suffix.lower() in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"):
            return _extract_image_ocr(path)
    except Exception as e:  # never let a bad file 500 the endpoint
        log.exception("Extraction failed for %s", path)
        return _empty(error=f"Unhandled extractor error: {e}")

    return _empty(error=f"Unsupported document type: {mime or path.suffix}")


def _empty(error: Optional[str] = None) -> Dict:
    return {
        "text": "",
        "page_count": None,
        "tables": [],
        "has_images": False,
        "ocr_applied": False,
        "error": error,
    }


# =====================================================================
# PDF
# =====================================================================

def _extract_pdf(path: Path) -> Dict:
    """Two-stage strategy:
       1. pypdf pulls the text layer.
       2. If the text layer is empty (scanned PDF) AND pytesseract is
          available, OCR every page.
       Tables come from pdfplumber when installed; otherwise skipped.
    """

    result = _empty()
    text_parts: List[str] = []
    page_count = 0
    partial_errors: List[str] = []

    # --- 1. pypdf text layer ---
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        for i, page in enumerate(reader.pages):
            try:
                text_parts.append(page.extract_text() or "")
            except Exception as e:
                partial_errors.append(f"page {i + 1}: {e}")
    except ImportError:
        partial_errors.append("pypdf not installed")
    except Exception as e:
        partial_errors.append(f"pypdf: {e}")

    result["page_count"] = page_count or None

    joined = "\n\n".join(t.strip() for t in text_parts if t and t.strip())

    # --- 2. If text layer empty, OCR ---
    if not joined.strip():
        ocr_text, ocr_ok, ocr_err = _ocr_pdf(path)
        if ocr_ok:
            joined = ocr_text
            result["ocr_applied"] = True
            result["has_images"] = True
        elif ocr_err:
            partial_errors.append(ocr_err)

    result["text"] = joined

    # --- 3. Tables via pdfplumber (best effort) ---
    try:
        import pdfplumber
        tables: List[Dict] = []
        with pdfplumber.open(str(path)) as pdf:
            for page_idx, page in enumerate(pdf.pages, start=1):
                for tbl in (page.extract_tables() or []):
                    if not tbl or len(tbl) < 2:
                        continue
                    header = [str(c or "").strip() for c in tbl[0]]
                    rows = [[str(c or "").strip() for c in r] for r in tbl[1:]]
                    tables.append({
                        "sheet":   None,
                        "page":    page_idx,
                        "columns": header,
                        "rows":    rows,
                    })
        result["tables"] = tables
    except ImportError:
        pass    # tables are a bonus, not a hard requirement
    except Exception as e:
        partial_errors.append(f"pdfplumber: {e}")

    if partial_errors:
        result["error"] = " | ".join(partial_errors)

    return result


def _ocr_pdf(path: Path) -> tuple:
    """Rasterise every PDF page and OCR it. Returns (text, ok, error)."""

    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        return "", False, f"OCR unavailable: {e}"

    try:
        pages = convert_from_path(str(path), dpi=200)
    except Exception as e:
        return "", False, f"pdf2image: {e} (poppler-utils may be missing)"

    parts = []
    for i, img in enumerate(pages, start=1):
        try:
            parts.append(pytesseract.image_to_string(img, lang="eng+tam+mal+hin"))
        except pytesseract.TesseractNotFoundError:
            return "", False, "tesseract binary not installed"
        except Exception as e:
            parts.append(f"[OCR error page {i}: {e}]")
    return "\n\n".join(p for p in parts if p and p.strip()), True, None


# =====================================================================
# DOCX
# =====================================================================

def _extract_docx(path: Path) -> Dict:
    result = _empty()
    try:
        import docx
    except ImportError:
        result["error"] = "python-docx not installed"
        return result

    try:
        doc = docx.Document(str(path))
    except Exception as e:
        result["error"] = f"docx open failed: {e}"
        return result

    # Paragraphs (skipping empty lines)
    para_text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

    # Tables
    tables: List[Dict] = []
    for t_idx, table in enumerate(doc.tables):
        rows_data = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
        ]
        if not rows_data:
            continue
        header = rows_data[0]
        body = rows_data[1:] if len(rows_data) > 1 else []
        tables.append({
            "sheet":   None,
            "page":    None,
            "columns": header,
            "rows":    body,
        })

    # Images — python-docx exposes them via inline_shapes / rels
    has_images = any("image" in (r.target_ref or "") for r in doc.part.rels.values())

    result["text"]       = para_text
    result["tables"]     = tables
    result["has_images"] = has_images
    return result


# =====================================================================
# XLSX
# =====================================================================

def _extract_xlsx(path: Path) -> Dict:
    result = _empty()
    try:
        from openpyxl import load_workbook
    except ImportError:
        result["error"] = "openpyxl not installed"
        return result

    try:
        wb = load_workbook(str(path), data_only=True, read_only=True)
    except Exception as e:
        result["error"] = f"openpyxl open failed: {e}"
        return result

    tables: List[Dict] = []
    text_parts: List[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_iter = ws.iter_rows(values_only=True)
        rows_list = [list(r) for r in rows_iter]
        if not rows_list:
            continue

        # Drop fully-empty leading rows to find the header
        while rows_list and all(v is None or str(v).strip() == "" for v in rows_list[0]):
            rows_list.pop(0)
        if not rows_list:
            continue

        header = [(str(c).strip() if c is not None else "") for c in rows_list[0]]
        body   = [
            [(str(c).strip() if c is not None else "") for c in row]
            for row in rows_list[1:]
        ]

        tables.append({
            "sheet":   sheet_name,
            "page":    None,
            "columns": header,
            "rows":    body,
        })

        # Also flatten to text so RAG retrieval works on plain sentences
        text_parts.append(f"[Sheet: {sheet_name}]")
        text_parts.append(" | ".join(header))
        for r in body:
            text_parts.append(" | ".join(r))
        text_parts.append("")

    result["text"]   = "\n".join(text_parts)
    result["tables"] = tables
    return result


# =====================================================================
# CSV
# =====================================================================

def _extract_csv(path: Path) -> Dict:
    import csv
    result = _empty()
    text_parts: List[str] = []
    columns: List[str] = []
    rows: List[List[str]] = []
    try:
        with path.open("r", encoding="utf-8", errors="replace", newline="") as fh:
            reader = csv.reader(fh)
            for i, row in enumerate(reader):
                row = [(c or "").strip() for c in row]
                text_parts.append(" | ".join(row))
                if i == 0:
                    columns = row
                else:
                    rows.append(row)
    except Exception as e:
        result["error"] = f"csv read failed: {e}"
        return result

    result["text"] = "\n".join(text_parts)
    if columns and rows:
        result["tables"] = [{
            "sheet":   None,
            "page":    None,
            "columns": columns,
            "rows":    rows,
        }]
    return result


# =====================================================================
# Plain text / markdown
# =====================================================================

def _extract_text(path: Path) -> Dict:
    result = _empty()
    try:
        result["text"] = path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        result["error"] = f"text read failed: {e}"
    return result


# =====================================================================
# Standalone image OCR
# =====================================================================

def _extract_image_ocr(path: Path) -> Dict:
    result = _empty()
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        result["error"] = f"OCR unavailable: {e}"
        return result

    try:
        img = Image.open(str(path))
        text = pytesseract.image_to_string(img, lang="eng+tam+mal+hin")
        result["text"] = text
        result["ocr_applied"] = True
        result["has_images"] = True
    except pytesseract.TesseractNotFoundError:
        result["error"] = "tesseract binary not installed"
    except Exception as e:
        result["error"] = f"OCR failed: {e}"
    return result
