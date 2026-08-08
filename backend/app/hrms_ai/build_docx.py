"""
Convert docs/HRMS_KNOWLEDGE.md to docs/HRMS_KNOWLEDGE.docx.

Runs standalone (no server dependency). Used to hand a
professionally-formatted Word copy of the HRMS knowledge base to
non-technical stakeholders — the .md file is still the authoritative
source that feeds the AI's vector index.

    python -m app.hrms_ai.build_docx

Requires python-docx (already in requirements.txt).
"""

from __future__ import annotations

import re
from pathlib import Path


_HERE = Path(__file__).resolve().parent
PROJECT_ROOT = _HERE.parent.parent.parent
MD_PATH = PROJECT_ROOT / "docs" / "HRMS_KNOWLEDGE.md"
DOCX_PATH = PROJECT_ROOT / "docs" / "HRMS_KNOWLEDGE.docx"


def build() -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    if not MD_PATH.exists():
        raise FileNotFoundError(f"Source Markdown not found: {MD_PATH}")

    md = MD_PATH.read_text(encoding="utf-8")

    doc = Document()

    # ---- Title page ----
    title = doc.add_heading("BVC24 HRMS Knowledge Base", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    p = doc.add_paragraph(
        "This document is the sole knowledge source for the BVC24 "
        "HRMS AI Assistant. Every module the assistant discusses is "
        "described here — attendance rules, leave policy, payroll "
        "calculations, memo automation, announcements, notifications "
        "and more."
    )
    for r in p.runs:
        r.font.size = Pt(11)

    doc.add_paragraph()

    # ---- Body ----
    for line in md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
        elif line.strip() == "":
            doc.add_paragraph()
        elif line.startswith("```"):
            # skip code fence markers; body lines inside a fence just
            # render as monospaced-looking regular paragraphs. Not
            # worth adding full parser complexity for this handoff.
            continue
        else:
            # Naive bold/italic support: strip Markdown markers.
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            doc.add_paragraph(clean.strip())

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    return DOCX_PATH


if __name__ == "__main__":
    out = build()
    print(f"[hrms_ai] Wrote {out}")
