"""
Resume parser — extract structured candidate data from a resume file.

Supported formats: PDF, DOCX, TXT, RTF (limited).

Two-layer extraction:
  1. Text extraction (PDF / DOCX / TXT -> plain text)
  2. Structured parsing (plain text -> ParsedResume) via:
       a. Regex + keyword rules (fast, deterministic, no LLM needed)
       b. Optional Gemini pass that refines/fills the rule-based output
          when GEMINI_API_KEY is available

The rule layer alone is good enough for ~80% of resumes. The Gemini
layer brings it close to 100% for messy / non-standard resumes.

No paid dependencies — pypdf and python-docx are MIT/BSD-licensed.
"""

from __future__ import annotations

from dataclasses import dataclass, asdict, field
from typing import List, Optional, Dict, Any
import os
import re
import io
import json


# =====================================================================
# Skill vocabulary — used to recognise skills in free text. Extend as
# needed; this is intentionally domain-aware (mechanical + software +
# manufacturing because BVC24 hires in both stacks).
# =====================================================================

_SKILL_VOCAB = sorted({
    # Software
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go",
    "ruby", "php", "kotlin", "swift", "scala", "rust",
    "react", "react.js", "react native", "vue", "vue.js", "angular",
    "next.js", "nuxt", "node", "node.js", "express", "fastapi", "flask",
    "django", "spring", "spring boot", ".net", "asp.net", "laravel",
    "html", "html5", "css", "css3", "sass", "scss", "tailwind",
    "tailwind css", "bootstrap", "jquery", "redux",
    "mysql", "postgresql", "postgres", "mongodb", "redis", "sqlite",
    "oracle", "mssql", "sql server", "mariadb", "elasticsearch",
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "jenkins", "github actions", "gitlab ci", "terraform", "ansible",
    "linux", "ubuntu", "centos", "windows server",
    "rest api", "rest", "graphql", "grpc", "websocket", "kafka",
    "rabbitmq", "celery", "airflow",
    "selenium", "pytest", "junit", "jest", "cypress", "playwright",
    "git", "github", "gitlab", "bitbucket", "jira", "confluence",
    "ml", "machine learning", "tensorflow", "pytorch", "keras",
    "scikit-learn", "numpy", "pandas", "matplotlib", "opencv",
    "nlp", "computer vision", "deep learning", "llm", "rag",
    "data analysis", "power bi", "tableau", "excel", "vba",
    # Mechanical / Manufacturing
    "autocad", "solidworks", "catia", "ansys", "creo", "nx",
    "matlab", "labview", "plc", "scada", "hmi",
    "mechanical design", "cnc", "vmc", "lathe", "milling",
    "welding", "fabrication", "sheet metal", "casting", "machining",
    "tig", "mig", "arc welding", "soldering",
    "iso 9001", "six sigma", "lean manufacturing", "5s", "kaizen",
    "production planning", "quality control", "qc", "qa",
    "bom", "erp", "sap", "oracle erp", "tally",
    "instrumentation", "pneumatics", "hydraulics",
    "electrical", "electronics", "embedded", "iot", "arduino",
    "raspberry pi", "stm32", "pcb design", "altium", "eagle",
    # Soft / business
    "leadership", "project management", "agile", "scrum", "kanban",
    "communication", "negotiation", "presentation", "training",
    "customer service", "sales", "marketing", "seo", "content writing",
    "accounting", "tally", "gst", "payroll", "tds",
}, key=len, reverse=True)


_DEGREE_PATTERNS = [
    r"\bph\.?d\b",
    r"\bm\.?s\.?\b", r"\bms\b", r"\bmsc\b", r"\bm\.?sc\b",
    r"\bm\.?tech\b", r"\bm\.?e\.?\b", r"\bmba\b",
    r"\bmca\b", r"\bmcom\b", r"\bma\b",
    r"\bb\.?tech\b", r"\bb\.?e\.?\b", r"\bbsc\b", r"\bb\.?sc\b",
    r"\bbca\b", r"\bbcom\b", r"\bba\b",
    r"\bdiploma\b", r"\bpolytechnic\b",
    r"\biti\b",
    r"\b12th\b", r"\bhsc\b", r"\bsslc\b", r"\b10th\b",
]


_LANGUAGES_VOCAB = sorted({
    "english", "hindi", "tamil", "telugu", "kannada", "malayalam",
    "marathi", "bengali", "gujarati", "punjabi", "urdu",
    "spanish", "french", "german", "japanese", "mandarin", "arabic",
}, key=len, reverse=True)


# =====================================================================
# Public schema
# =====================================================================

@dataclass
class ParsedResume:
    full_name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None

    skills: List[str] = field(default_factory=list)
    languages: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)

    education: List[Dict[str, Any]] = field(default_factory=list)
    # [{"degree": "...", "institution": "...", "year": 2018}]

    work_experience: List[Dict[str, Any]] = field(default_factory=list)
    # [{"company": "...", "role": "...", "from": "...", "to": "..."}]

    projects: List[Dict[str, Any]] = field(default_factory=list)
    # [{"name": "...", "description": "..."}]

    total_experience_years: Optional[float] = None
    highest_qualification: Optional[str] = None

    raw_text: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


# =====================================================================
# Public entry point
# =====================================================================

def parse_resume(filename: str, file_bytes: bytes) -> ParsedResume:
    """
    Main entry point. Extracts text from `file_bytes` based on `filename`'s
    extension, then runs rule-based + (optional) LLM parsing.
    """
    text = _extract_text(filename, file_bytes)
    if not text or not text.strip():
        return ParsedResume(raw_text="")

    rule_based = _rule_based_parse(text)
    rule_based.raw_text = text

    # Optional LLM refinement — fills/corrects any fields the rule
    # layer couldn't handle. Best-effort; never throws.
    llm = _llm_refine(text, rule_based)
    if llm:
        rule_based = _merge(rule_based, llm)

    return rule_based


# =====================================================================
# Text extraction
# =====================================================================

def _extract_text(filename: str, file_bytes: bytes) -> str:
    ext = (filename or "").lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext in ("docx", "doc"):
        return _extract_docx(file_bytes)
    if ext in ("txt", "rtf"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    # Best-effort decode for unknown formats
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_pdf(file_bytes: bytes) -> str:
    """Try pypdf -> PyPDF2 -> pdfplumber, gracefully degrade if no
    library is installed."""
    # pypdf (modern)
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(
            (page.extract_text() or "") for page in reader.pages
        )
    except Exception:
        pass
    # legacy PyPDF2
    try:
        from PyPDF2 import PdfReader  # type: ignore
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(
            (page.extract_text() or "") for page in reader.pages
        )
    except Exception:
        pass
    # pdfplumber as a last resort
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            return "\n".join((p.extract_text() or "") for p in pdf.pages)
    except Exception:
        pass
    return ""


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx  # type: ignore  -- from python-docx
        f = io.BytesIO(file_bytes)
        d = docx.Document(f)
        parts: List[str] = []
        for p in d.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text)
        # also pull table cell text
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


# =====================================================================
# Rule-based parsing
# =====================================================================

_EMAIL_RE     = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RE     = re.compile(r"(?:\+?\d{1,3}[\s\-]?)?(?:\(?\d{3,5}\)?[\s\-]?)?\d{3,5}[\s\-]?\d{3,6}")
_LINKEDIN_RE  = re.compile(r"linkedin\.com/in/[A-Za-z0-9_\-]+", re.IGNORECASE)
_YEAR_RE      = re.compile(r"(?:19|20)\d{2}")
_INDIAN_CITIES = sorted({
    "chennai", "coimbatore", "bangalore", "bengaluru", "hyderabad",
    "mumbai", "pune", "delhi", "noida", "gurgaon", "gurugram",
    "kolkata", "ahmedabad", "kochi", "thiruvananthapuram",
    "madurai", "trichy", "tiruchirappalli", "salem", "erode",
    "jaipur", "lucknow", "indore", "bhopal", "nagpur", "vadodara",
    "vizag", "visakhapatnam", "vijayawada",
}, key=len, reverse=True)


def _rule_based_parse(text: str) -> ParsedResume:
    out = ParsedResume()
    lower = text.lower()

    # Email
    m = _EMAIL_RE.search(text)
    if m:
        out.email = m.group(0)

    # Phone — try to find a plausible 10-digit Indian number first
    out.phone = _best_phone(text)

    # LinkedIn URL
    m = _LINKEDIN_RE.search(text)
    if m:
        out.linkedin = "https://" + m.group(0) if not m.group(0).startswith("http") else m.group(0)

    # Name — first non-empty line that looks like a person's name (2-4 words,
    # title-cased, no digits, no email keywords)
    for line in (l.strip() for l in text.splitlines()):
        if not line or "@" in line or any(c.isdigit() for c in line):
            continue
        words = line.split()
        if 2 <= len(words) <= 5 and sum(1 for w in words if w[:1].isupper()) >= len(words) - 1:
            if not any(k in line.lower() for k in [
                "resume", "curriculum", "cv", "phone", "email", "address",
                "contact", "objective", "summary",
            ]):
                out.full_name = line
                break

    # Location — first Indian city match
    for city in _INDIAN_CITIES:
        if re.search(rf"\b{re.escape(city)}\b", lower):
            out.location = city.title()
            break

    # Skills
    out.skills = _extract_skills(lower)

    # Languages
    out.languages = _extract_languages(lower)

    # Certifications — line-based heuristic
    out.certifications = _extract_certifications(text)

    # Education
    out.education = _extract_education(text)
    if out.education:
        out.highest_qualification = _highest_qualification(out.education)

    # Work experience
    out.work_experience = _extract_work_experience(text)

    # Total experience years (best-effort estimate)
    out.total_experience_years = _estimate_experience_years(text, out.work_experience)

    return out


def _best_phone(text: str) -> Optional[str]:
    candidates = []
    for m in _PHONE_RE.finditer(text):
        s = re.sub(r"[\s\-()]", "", m.group(0))
        # Strip leading 0 / +91
        s = re.sub(r"^(\+?91)?0?", "", s)
        if 10 <= len(s) <= 12:
            candidates.append(s)
    # Prefer the first 10-digit one (Indian mobile)
    for c in candidates:
        if len(c) == 10 and c[0] in "6789":
            return c
    return candidates[0] if candidates else None


def _extract_skills(lower_text: str) -> List[str]:
    found = set()
    for s in _SKILL_VOCAB:
        # Boundary match to avoid "ms" hitting "msexcel" etc.
        if re.search(rf"(?<![A-Za-z0-9]){re.escape(s)}(?![A-Za-z0-9])", lower_text):
            found.add(s)
    return sorted(found, key=lambda x: (-len(x), x))


def _extract_languages(lower_text: str) -> List[str]:
    found = set()
    for l in _LANGUAGES_VOCAB:
        if re.search(rf"\b{re.escape(l)}\b", lower_text):
            found.add(l.title())
    return sorted(found)


def _extract_certifications(text: str) -> List[str]:
    """Heuristic — look inside a 'Certifications' section. Cap at 10 lines."""
    out: List[str] = []
    lines = text.splitlines()
    in_section = False
    for line in lines:
        s = line.strip()
        if not s:
            continue
        s_lower = s.lower()
        if not in_section:
            if any(h in s_lower for h in ["certification", "certificate"]) and len(s) < 60:
                in_section = True
            continue
        # End section on next major header
        if any(h in s_lower for h in [
            "education", "experience", "projects", "skills",
            "achievements", "personal details", "languages",
        ]):
            break
        if s.startswith(("-", "*", "•")) or len(s) > 6:
            out.append(s.lstrip("-*• ").strip())
        if len(out) >= 10:
            break
    return out


def _extract_education(text: str) -> List[Dict[str, Any]]:
    """Find degree-keyword lines and pull the surrounding context."""
    edu: List[Dict[str, Any]] = []
    for ln in text.splitlines():
        s = ln.strip()
        s_lower = s.lower()
        if not s:
            continue
        for pat in _DEGREE_PATTERNS:
            if re.search(pat, s_lower):
                year = None
                m = _YEAR_RE.search(s)
                if m:
                    try:
                        year = int(m.group(0))
                    except ValueError:
                        pass
                edu.append({
                    "degree": s[:160],
                    "year": year,
                })
                break
    # Dedupe by degree text
    seen, dedup = set(), []
    for e in edu:
        key = e["degree"].lower()
        if key not in seen:
            seen.add(key)
            dedup.append(e)
    return dedup[:8]


_DEGREE_RANK = [
    ("phd", 8), ("ph.d", 8), ("doctor", 8),
    ("m.tech", 7), ("mtech", 7), ("m.e.", 7), ("me ", 7),
    ("ms ", 6), ("m.sc", 6), ("msc", 6), ("mba", 6), ("mca", 6),
    ("b.tech", 5), ("btech", 5), ("b.e.", 5), ("be ", 5),
    ("bsc", 4), ("b.sc", 4), ("bca", 4), ("bcom", 4), ("ba ", 4),
    ("diploma", 3), ("polytechnic", 3),
    ("iti", 2),
    ("12th", 1), ("hsc", 1), ("10th", 1), ("sslc", 1),
]


def _highest_qualification(edu: List[Dict[str, Any]]) -> Optional[str]:
    best_rank, best_str = -1, None
    for e in edu:
        d = (e.get("degree") or "").lower()
        for key, rank in _DEGREE_RANK:
            if key in d and rank > best_rank:
                best_rank, best_str = rank, e["degree"]
    return best_str


_COMPANY_HINTS = re.compile(
    r"\b(at|@)\s+([A-Z][A-Za-z0-9.&\-, ]{2,80}?)\b(?:\s+from|\s+\(|\s+\d{4})?",
)


def _extract_work_experience(text: str) -> List[Dict[str, Any]]:
    """Heuristic: pull 'at <Company> from <year> to <year>' patterns plus
    lines under an 'Experience' header that mention years."""
    out: List[Dict[str, Any]] = []

    # 1. Year-range pattern  (2019 - 2022 / 2019 to Present / 2019-2024)
    range_re = re.compile(
        r"(?P<from>(?:19|20)\d{2})\s*(?:-|to|–|—)\s*"
        r"(?P<to>(?:19|20)\d{2}|present|current|now)",
        re.IGNORECASE,
    )

    for ln in text.splitlines():
        m = range_re.search(ln)
        if not m:
            continue
        out.append({
            "from": m.group("from"),
            "to":   m.group("to"),
            "role_company": ln.strip()[:200],
        })

    return out[:10]


def _estimate_experience_years(
    text: str, work_experience: List[Dict[str, Any]]
) -> Optional[float]:
    """Several strategies; pick the first that produces a sensible number."""
    # 1. Explicit "X years of experience" in text
    m = re.search(
        r"(\d{1,2}(?:\.\d{1,2})?)\s*\+?\s*years?\s+(?:of\s+)?experience",
        text.lower(),
    )
    if m:
        try:
            n = float(m.group(1))
            if 0 < n < 50:
                return n
        except ValueError:
            pass

    # 2. Sum year-range entries
    total = 0.0
    for we in work_experience:
        try:
            frm = int(we["from"])
            to_str = (we["to"] or "").lower()
            if to_str in ("present", "current", "now"):
                from datetime import date
                to = date.today().year
            else:
                to = int(to_str)
            if frm < to <= 2100:
                total += (to - frm)
        except Exception:
            continue
    if total > 0:
        return round(min(total, 50), 1)

    return None


# =====================================================================
# Optional Gemini refinement
# =====================================================================

def _llm_refine(text: str, rule_based: ParsedResume) -> Optional[Dict[str, Any]]:
    """Ask Gemini to extract anything the rules missed. Returns a dict
    of additive fields. Best-effort — returns None on any failure."""
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if not api_key:
        return None

    try:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
    except Exception:
        return None

    snippet = text[:6000]   # bound the prompt; resumes rarely need more
    have = rule_based.to_dict()
    have.pop("raw_text", None)
    have_str = json.dumps(have, default=str)

    prompt = (
        "You are a resume parser. Extract structured candidate data from "
        "the resume text below. Return STRICT JSON only — no markdown.\n\n"
        "Schema:\n"
        "{ \"full_name\": str|null, \"email\": str|null, \"phone\": str|null, "
        "\"location\": str|null, \"linkedin\": str|null, "
        "\"skills\": [str], \"languages\": [str], \"certifications\": [str], "
        "\"education\": [{\"degree\": str, \"institution\": str|null, \"year\": int|null}], "
        "\"work_experience\": [{\"company\": str|null, \"role\": str|null, "
        "\"from\": str|null, \"to\": str|null, \"description\": str|null}], "
        "\"projects\": [{\"name\": str, \"description\": str|null}], "
        "\"total_experience_years\": number|null, "
        "\"highest_qualification\": str|null }\n\n"
        "Rules:\n"
        " - Fill ONLY fields you can find evidence for.\n"
        " - Skills must be tech / domain skills, NOT soft-skill fluff.\n"
        " - Do not invent companies, dates, or degrees.\n"
        " - 'total_experience_years' = numeric, e.g. 4.5\n\n"
        f"RULE-BASED OUTPUT SO FAR (you may improve or extend):\n{have_str}\n\n"
        f"RESUME TEXT:\n{snippet}\n\n"
        "Return JSON only:"
    )

    for model_name in ["gemini-2.5-flash-lite", "gemini-2.0-flash-lite",
                      "gemini-2.5-flash", "gemini-2.0-flash"]:
        try:
            model = genai.GenerativeModel(
                model_name,
                generation_config={
                    "temperature": 0.2,
                    "max_output_tokens": 2000,
                    "response_mime_type": "application/json",
                },
            )
            resp = model.generate_content(prompt)
            parsed = _safe_json(resp.text or "")
            if parsed:
                return parsed
        except Exception:
            continue
    return None


def _safe_json(s: str) -> Optional[Dict[str, Any]]:
    if not s:
        return None
    s = s.strip()
    if s.startswith("```"):
        s = re.sub(r"^```[a-zA-Z]*\s*", "", s)
        s = re.sub(r"\s*```$", "", s)
    try:
        return json.loads(s)
    except Exception:
        m = re.search(r"\{.*\}", s, re.DOTALL)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                return None
    return None


def _merge(base: ParsedResume, refinements: Dict[str, Any]) -> ParsedResume:
    """Overlay LLM refinements on the rule-based output, preferring the
    LLM value when present and non-empty (except for raw_text)."""
    out = ParsedResume(**asdict(base))

    def pick(field_name, current):
        v = refinements.get(field_name)
        if isinstance(v, str) and v.strip():
            return v.strip()
        if isinstance(v, (int, float)) and v:
            return v
        if isinstance(v, list) and v:
            return v
        return current

    out.full_name             = pick("full_name", out.full_name)
    out.email                 = pick("email", out.email)
    out.phone                 = pick("phone", out.phone)
    out.location              = pick("location", out.location)
    out.linkedin              = pick("linkedin", out.linkedin)
    out.skills                = pick("skills", out.skills)
    out.languages             = pick("languages", out.languages)
    out.certifications        = pick("certifications", out.certifications)
    out.education             = pick("education", out.education)
    out.work_experience       = pick("work_experience", out.work_experience)
    out.projects              = pick("projects", out.projects)
    out.total_experience_years = pick("total_experience_years", out.total_experience_years)
    out.highest_qualification = pick("highest_qualification", out.highest_qualification)
    return out
