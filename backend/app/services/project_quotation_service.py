"""
Project Quotation Template rendering — builds the default CONTENT_JSON for
a brand-new project, compiles it (+ live/overridden company info) into a
full HTML document for preview, and renders that into PDF (xhtml2pdf) and
Word (python-docx) bytes.

Plain functions only (no classes), matching every other service module in
this codebase. Never raise on rendering failure — return None so the
calling route decides how to surface it (matches quotation.py's own
render_quotation_pdf, which returns None on failure).
"""

import json
import logging
import re
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Optional

from app.services.company_settings_service import format_full_address, serialize_company

log = logging.getLogger(__name__)

_STATIC_DIR = Path(__file__).resolve().parent.parent.parent / "static"


def _slugify(value: str) -> str:
    value = (value or "").strip().lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "project"


def _financial_year(d: date) -> str:
    """Indian financial year label, e.g. 24-25 for any date between
    2024-04-01 and 2025-03-31."""
    start_year = d.year if d.month >= 4 else d.year - 1
    return f"{str(start_year)[-2:]}-{str(start_year + 1)[-2:]}"


def default_quotation_number(project, company, qdate: date) -> str:
    prefix = (company.SHORT_NAME or company.LEGAL_NAME or "CO").split()[0].upper()
    return f"{prefix}/QTN/{_financial_year(qdate)}/{_slugify(project.NAME)}"


def build_default_quotation_content(project, company, pricing=None) -> dict:
    """Pure — no DB/IO. Seeds a new project's quotation, loosely mirroring
    the sample PDF's structure (red banner header, intro, price schedule,
    16-point Terms & Conditions, footer).

    `pricing` (an optional ProjectPricing row) seeds the price-schedule row's
    default unit price from ORIGINAL_PRICE when present; falls back to 0
    otherwise — unchanged from before pricing existed."""
    bank_lines = []
    if company.BANK_NAME:
        bank_lines.append(f"Bank: {company.BANK_NAME}")
    if company.BANK_ACCOUNT_NUMBER:
        bank_lines.append(f"Acc. No: {company.BANK_ACCOUNT_NUMBER}")
    if company.BANK_IFSC:
        bank_lines.append(f"IFSC: {company.BANK_IFSC}")
    if company.BANK_BRANCH:
        bank_lines.append(f"Branch: {company.BANK_BRANCH}")
    if company.UPI_ID:
        bank_lines.append(f"UPI: {company.UPI_ID}")
    bank_details = " | ".join(bank_lines) or "To be provided"

    terms_items = [
        "Price : Ex-works.",
        "Order : Purchase Order to be placed in the company's name.",
        "Delivery : As mutually agreed upon confirmation of order with advance payment.",
        "Change of Specification : We reserve the right to change/modify the specification necessary during design and development for betterment/improvement without prior notice.",
        "Inspection : To be carried out at our works prior to dispatch, subject to acceptance by the visiting authority.",
        "Local Taxes : GST as applicable, extra.",
        "Packing &amp; Handling : Extra, at actuals.",
        "Branding : At actual, to your account.",
        "Freight : At actual, to your account.",
        "Insurance : At actual, to your account.",
        "Payment Terms : As mutually agreed, along with Purchase Order.",
        "Installation Charges : Installation and commissioning charges, if applicable, to your account.",
        "Validity : 15 days from the quotation date.",
        "Warranty : Manufacturing defects covered for 12 months from the date of supply.",
        "Cancellation : Once placed, an order cannot be cancelled. Advance payment will be adjusted against expenses incurred till then.",
        f"Bank Details : {bank_details}",
    ]
    terms_html = "<ol>" + "".join(f"<li>{t}</li>" for t in terms_items) + "</ol>"

    return {
        "version": 1,
        "letterhead": {
            "showCompanyLogo": True,
            "logoUrlOverride": None,
            "bannerColor": "#C8102E",
            "title": "QUOTATION",
            "subtitle": project.NAME,
        },
        "companyInfoOverride": None,
        "customerInfo": {
            "showBlock": True,
            "label": "BILL TO",
            "namePlaceholder": "{{customer_name}}",
            "addressPlaceholder": "{{customer_address}}",
            "freeText": "",
        },
        "introHtml": f"<p>Dear Sir/Madam,</p><p>We are pleased to submit our quotation for <strong>{project.NAME}</strong> for your kind consideration.{(' ' + project.DESCRIPTION) if project.DESCRIPTION else ''}</p>",
        "sections": [
            {
                "id": "table-1",
                "type": "table",
                "order": 0,
                "title": "Price Schedule",
                "rows": [
                    {"description": project.NAME + (f"\n{project.DESCRIPTION}" if project.DESCRIPTION else ""),
                     "qty": 1, "unitPrice": float(pricing.ORIGINAL_PRICE) if pricing is not None else 0},
                ],
                "pageBreakBefore": False,
            },
        ],
        "termsHtml": terms_html,
        "notesHtml": "",
        "footerHtml": "<p>{{company_website}} &middot; {{company_address}}</p>",
        "signature": {"imageUrl": None, "name": "", "designation": "", "showBlock": True},
        "style": {"accentColor": "#C8102E", "pageMargin": "18mm 14mm"},
    }


def _resolve_company_info(content: dict, company) -> dict:
    """Override (frozen per-quotation snapshot) if set, else live CompanyMaster —
    preserves Project Isolation (an override never touches the shared row or
    any other project's quotation)."""
    override = content.get("companyInfoOverride")
    if override:
        return override
    info = serialize_company(company)
    info["FULL_ADDRESS"] = format_full_address(company)
    return info


def _substitute_tokens(html: str, tokens: dict) -> str:
    for key, value in tokens.items():
        html = html.replace("{{" + key + "}}", str(value) if value is not None else "")
    return html


def render_quotation_html(quotation_row, company) -> str:
    """Merges CONTENT_JSON + live-or-overridden company info into a full
    HTML document, reusing the same @page/print-CSS conventions already
    established in quotation.py's _build_quotation_pdf_html() so output
    visually matches the rest of the app's document family."""
    content = json.loads(quotation_row.CONTENT_JSON)
    info = _resolve_company_info(content, company)
    company_addr = info.get("FULL_ADDRESS") or format_full_address(company)

    tokens = {
        "customer_name": "",
        "customer_address": "",
        "company_legal_name": info.get("LEGAL_NAME") or "",
        "company_address": company_addr,
        "company_website": info.get("WEBSITE") or "",
    }

    letterhead = content.get("letterhead", {})
    style = content.get("style", {})
    accent = style.get("accentColor", "#C8102E")
    margin = style.get("pageMargin", "18mm 14mm")

    logo_url = letterhead.get("logoUrlOverride") or info.get("LOGO_URL")
    # xhtml2pdf does not reliably honor CSS max-height/max-width on <img> —
    # it falls back to the image's native pixel size, which can dwarf the
    # page for a large upload. An explicit height (width scales
    # proportionally) is the combination xhtml2pdf actually respects.
    logo_html = f'<img src="{logo_url}" style="height:48px; margin-bottom:8px;"/>' if logo_url else ""

    customer = content.get("customerInfo", {})
    customer_html = ""
    if customer.get("showBlock", True):
        customer_html = f"""
        <div class="label">{customer.get('label', 'BILL TO')}</div>
        <div style="font-weight:bold;">{customer.get('namePlaceholder', '')}</div>
        <div>{customer.get('addressPlaceholder', '')}</div>
        {(f"<div>{customer.get('freeText')}</div>") if customer.get('freeText') else ''}
        """

    sections_html = ""
    for section in sorted(content.get("sections", []), key=lambda s: s.get("order", 0)):
        page_break = '<div style="page-break-before: always;"></div>' if section.get("pageBreakBefore") else ""
        title_html = f'<div class="section-title">{section["title"]}</div>' if section.get("title") else ""
        stype = section.get("type")
        if stype == "richtext" or stype == "custom":
            body = f'<div class="section-body">{section.get("html", "")}</div>'
        elif stype == "table":
            rows = section.get("rows", [])
            rows_html = ""
            for idx, r in enumerate(rows, start=1):
                qty = r.get("qty") or 0
                unit_price = r.get("unitPrice") or 0
                amount = float(qty) * float(unit_price)
                desc = (r.get("description") or "").replace("\n", "<br/>")
                rows_html += f"""
                <tr>
                  <td style="text-align:center;">{idx}</td>
                  <td>{desc}</td>
                  <td style="text-align:right;">{qty}</td>
                  <td style="text-align:right;">{unit_price:,.2f}</td>
                  <td style="text-align:right; font-weight:bold;">{amount:,.2f}</td>
                </tr>
                """
            body = f"""
            <table class="items">
              <thead><tr>
                <th style="width:30px; text-align:center;">#</th>
                <th>Product Description</th>
                <th style="width:60px; text-align:right;">Qty</th>
                <th style="width:100px; text-align:right;">Unit Price</th>
                <th style="width:100px; text-align:right;">Amount</th>
              </tr></thead>
              <tbody>{rows_html or '<tr><td colspan="5" style="text-align:center; color:#94a3b8;">No items</td></tr>'}</tbody>
            </table>
            """
        elif stype == "image":
            # Same xhtml2pdf max-width caveat as the logo/signature images above.
            body = f'<img src="{section.get("imageUrl", "")}" style="width:100%;"/>'
        else:
            body = ""
        sections_html += page_break + title_html + body

    signature = content.get("signature", {})
    signature_html = ""
    if signature.get("showBlock", True) and (signature.get("name") or signature.get("imageUrl")):
        sig_img = f'<img src="{signature["imageUrl"]}" style="height:60px;"/>' if signature.get("imageUrl") else ""
        signature_html = f"""
        <div class="signature">
          <p>Thanks &amp; Regards,</p>
          <p>For {info.get('LEGAL_NAME') or ''},</p>
          {sig_img}
          <p style="font-weight:bold; margin-top:6px;">{signature.get('name', '')}</p>
          <p>{signature.get('designation', '')}</p>
        </div>
        """

    intro_html = _substitute_tokens(content.get("introHtml", ""), tokens)
    terms_html = _substitute_tokens(content.get("termsHtml", ""), tokens)
    notes_html = _substitute_tokens(content.get("notesHtml", ""), tokens)
    footer_html = _substitute_tokens(content.get("footerHtml", ""), tokens)

    return f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<title>Quotation {quotation_row.QUOTATION_NUMBER}</title>
<style>
  @page {{ size: A4; margin: {margin}; }}
  body {{ font-family: Helvetica, Arial, sans-serif; font-size: 10pt; color: #0f172a; }}
  h1, h2, h3 {{ margin: 0; }}
  .header {{ background-color: {accent}; color: white; padding: 14px 18px; }}
  .header h1 {{ font-size: 18pt; }}
  .header .num {{ font-size: 9pt; opacity: 0.9; }}
  table {{ width: 100%; border-collapse: collapse; }}
  .meta td {{ padding: 4px 8px; vertical-align: top; }}
  .label {{ color: #64748b; font-size: 8pt; }}
  .items th {{ background-color: #fef2f2; color: #8B0B1F; padding: 8px; text-align: left; font-size: 9pt; border-bottom: 1px solid #fecaca; }}
  .items td {{ padding: 6px 8px; border-bottom: 1px solid #e2e8f0; }}
  .section-title {{ font-weight: bold; color: {accent}; margin: 14px 0 6px; }}
  .section-body {{ font-size: 9.5pt; line-height: 1.5; }}
  .terms {{ margin-top: 16px; background-color: #f8fafc; padding: 10px 12px; border-left: 3px solid {accent}; font-size: 9pt; }}
  .terms ol {{ margin: 4px 0 0 18px; padding: 0; }}
  .signature {{ margin-top: 24px; font-size: 9.5pt; }}
  .footer {{ margin-top: 20px; padding-top: 10px; border-top: 1px solid #e2e8f0; font-size: 8pt; color: #94a3b8; text-align: center; }}
</style>
</head>
<body>
  <div class="header">
    {logo_html}
    <h1>{letterhead.get('title', 'QUOTATION')}</h1>
    <div class="num">{letterhead.get('subtitle', '')}</div>
  </div>

  <table class="meta" style="margin-top: 14px;">
    <tr>
      <td style="width: 50%;">{customer_html}</td>
      <td style="width: 50%; text-align: right;">
        <div class="label">QUOTE REFERENCE NUMBER</div>
        <div style="font-weight: bold;">{quotation_row.QUOTATION_NUMBER}</div>
        <div class="label" style="margin-top: 6px;">QUOTE DATE</div>
        <div>{quotation_row.QUOTATION_DATE}</div>
      </td>
    </tr>
  </table>

  <div class="section-body" style="margin-top: 14px;">{intro_html}</div>

  {sections_html}

  {(f'<div class="terms"><div class="section-title" style="margin-top:0;">Terms &amp; Conditions</div>{terms_html}</div>') if terms_html else ''}

  {(f'<div class="section-body" style="margin-top:14px;">{notes_html}</div>') if notes_html else ''}

  {signature_html}

  <div class="footer">{footer_html}</div>
</body>
</html>
"""


def sync_final_price_into_quotation(quotation_row, pricing, company) -> bool:
    """Pushes ProjectPricing.FINAL_PRICE into the first table-type section's
    first row's unitPrice whenever pricing is saved, so the quotation's
    Amount column reflects the latest pricing automatically — without
    touching any other quotation content (other rows, sections, terms,
    letterhead, etc. are left exactly as the user left them). Re-renders
    RENDERED_HTML so the preview/PDF/DOCX pick up the change immediately.

    Returns True if a sync was applied (caller should commit), False if
    there was no table section to sync into (nothing to do — e.g. the user
    deleted the price-schedule table entirely)."""
    try:
        content = json.loads(quotation_row.CONTENT_JSON)
    except Exception:
        return False

    for section in content.get("sections", []):
        if section.get("type") == "table" and section.get("rows"):
            section["rows"][0]["unitPrice"] = float(pricing.FINAL_PRICE)
            quotation_row.CONTENT_JSON = json.dumps(content)
            quotation_row.RENDERED_HTML = render_quotation_html(quotation_row, company)
            return True
    return False


def _resolve_static_path(url: str) -> Optional[Path]:
    """Maps a stored '/static/...' URL (relative, browser-facing) to the
    actual file on disk. Neither xhtml2pdf nor python-docx can resolve a
    relative URL against the app's static mount on their own — both need
    a real filesystem path handed to them explicitly."""
    if not url or url.startswith("http://") or url.startswith("https://"):
        return None
    rel = url.lstrip("/")
    if rel.startswith("static/"):
        rel = rel[len("static/"):]
    path = _STATIC_DIR / rel
    return path if path.exists() else None


def render_quotation_pdf_bytes(html: str) -> tuple[Optional[bytes], Optional[str]]:
    """Thin wrapper around xhtml2pdf.pisa.CreatePDF, mirroring
    quotation.py's existing render_quotation_pdf. Returns (None, error_message)
    on any failure — the caller decides how to surface it. The error message
    is returned (not just logged) so the API response can show the real
    cause instead of a generic string."""
    try:
        from xhtml2pdf import pisa
    except ImportError as e:
        log.warning("xhtml2pdf not installed — cannot render quotation PDF")
        return None, f"xhtml2pdf is not installed: {e}"

    def _link_callback(uri, _rel):
        # xhtml2pdf hands us the raw <img src="..."> value; our HTML always
        # uses the app's relative "/static/..." URLs, which xhtml2pdf cannot
        # resolve on its own (it has no notion of the FastAPI static mount).
        # Map it to the real file on disk so images actually embed.
        resolved = _resolve_static_path(uri)
        return str(resolved) if resolved else uri

    buf = BytesIO()
    try:
        result = pisa.CreatePDF(html, dest=buf, link_callback=_link_callback)
        if result.err:
            log.warning("xhtml2pdf reported %s error(s) rendering quotation PDF", result.err)
            return None, f"xhtml2pdf reported {result.err} rendering error(s) — check server logs for details"
        return buf.getvalue(), None
    except Exception as e:
        log.exception("Quotation PDF render failed")
        return None, f"{type(e).__name__}: {e}"


# Image formats python-docx's own docx.image package recognizes natively.
# Anything else (e.g. WebP) must be converted to PNG in-memory via Pillow
# before add_picture() will accept it — otherwise it raises
# UnrecognizedImageError, which the caller would otherwise swallow silently.
_DOCX_NATIVE_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff"}


def _docx_image_source(path: Path):
    """Returns something add_picture() can embed for the given image file:
    the path itself if python-docx already recognizes the format, an
    in-memory PNG conversion if Pillow can read it but python-docx can't
    (WebP being the practical case here), or None if it can't be read at
    all (e.g. SVG, which is vector — Pillow cannot rasterize it)."""
    if path.suffix.lower() in _DOCX_NATIVE_IMAGE_EXTS:
        return str(path)
    try:
        from PIL import Image
    except ImportError:
        log.warning("Pillow not installed — cannot convert %s for DOCX embedding", path)
        return None
    try:
        with Image.open(path) as im:
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGBA" if "transparency" in im.info or im.mode == "P" else "RGB")
            buf = BytesIO()
            im.save(buf, format="PNG")
            buf.seek(0)
            return buf
    except Exception:
        log.warning("Could not convert image %s for DOCX embedding", path)
        return None


def render_quotation_docx_bytes(quotation_row, company) -> Optional[bytes]:
    """python-docx render, built to mirror render_quotation_html's structure
    and styling section-by-section (banner header, meta table, section
    bodies, items table, terms box, footer) so the Word output stays
    visually consistent with the PDF/preview instead of flattening to plain
    text. python-docx has no HTML importer, so a small HTML-to-paragraph
    renderer below handles the limited tag vocabulary RichTextEditor
    produces (bold/italic/underline/strike/lists/headings/line breaks) —
    "good enough" fidelity, not pixel-parity, but far closer than raw
    tag-stripping."""
    try:
        import docx
        from docx.shared import Pt, Inches, Mm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from html.parser import HTMLParser
    except ImportError:
        log.warning("python-docx not installed — cannot render quotation DOCX")
        return None

    def _rgb(hex_color, fallback="000000"):
        try:
            return RGBColor.from_string((hex_color or fallback).lstrip("#"))
        except Exception:
            return RGBColor.from_string(fallback)

    def _set_cell_shading(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        tcPr.append(shd)

    def _set_cell_border(cell, **edges):
        """edges: e.g. bottom={'sz':6,'color':'E2E8F0'}, left={'sz':24,'color':'C8102E'}."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge, spec in edges.items():
            if spec is None:
                continue
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), spec.get("val", "single"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:color"), spec.get("color", "000000"))
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _no_table_borders(table):
        for row in table.rows:
            for cell in row.cells:
                _set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"},
                                  left={"val": "nil"}, right={"val": "nil"})

    class _SimpleHtmlToDocx(HTMLParser):
        """Renders the small, controlled HTML vocabulary RichTextEditor emits
        (p/strong/b/em/i/u/s/strike/ul/ol/li/br/h1/h2/h3) into paragraphs on
        the given container (a Document, table cell, or footer — anything
        exposing add_paragraph()). Not a general HTML renderer."""

        _HEADING_PT = {"h1": 14, "h2": 12, "h3": 11}

        def __init__(self, container, accent_rgb=None, reuse_para=None):
            super().__init__(convert_charrefs=True)
            self.container = container
            self.accent_rgb = accent_rgb
            self.bold = self.italic = self.underline = self.strike = False
            self.heading = None
            self.list_type = None
            self.para = None
            self.wrote_anything = False
            # An already-existing empty paragraph (e.g. a footer's default
            # first paragraph) to reuse for the first block instead of
            # leaving it as a stray blank line before a brand-new one.
            self._reuse_para = reuse_para

        def _new_para(self, style=None):
            if self._reuse_para is not None:
                self.para = self._reuse_para
                self._reuse_para = None
                if style:
                    try:
                        self.para.style = style
                    except Exception:
                        pass
                return self.para
            try:
                self.para = self.container.add_paragraph(style=style)
            except KeyError:
                self.para = self.container.add_paragraph()
            return self.para

        def handle_starttag(self, tag, attrs):
            if tag == "p":
                self._new_para()
            elif tag in ("strong", "b"):
                self.bold = True
            elif tag in ("em", "i"):
                self.italic = True
            elif tag == "u":
                self.underline = True
            elif tag in ("s", "strike", "del"):
                self.strike = True
            elif tag in ("h1", "h2", "h3"):
                self.heading = tag
                self._new_para()
            elif tag in ("ul", "ol"):
                self.list_type = tag
            elif tag == "li":
                self._new_para(style="List Number" if self.list_type == "ol" else "List Bullet")
            elif tag == "br" and self.para is not None:
                self.para.add_run().add_break()

        def handle_endtag(self, tag):
            if tag in ("strong", "b"):
                self.bold = False
            elif tag in ("em", "i"):
                self.italic = False
            elif tag == "u":
                self.underline = False
            elif tag in ("s", "strike", "del"):
                self.strike = False
            elif tag in ("h1", "h2", "h3"):
                self.heading = None
            elif tag in ("ul", "ol"):
                self.list_type = None

        def handle_data(self, data):
            if self.para is None:
                if not data.strip():
                    return
                self._new_para()
            if not data:
                return
            run = self.para.add_run(data)
            run.bold = self.bold or bool(self.heading)
            run.italic = self.italic
            run.underline = self.underline
            if self.strike:
                run.font.strike = True
            if self.heading:
                run.font.size = Pt(self._HEADING_PT[self.heading])
                if self.accent_rgb:
                    run.font.color.rgb = self.accent_rgb
            self.wrote_anything = True

    def _render_html(container, html, accent_rgb=None, reuse_para=None):
        renderer = _SimpleHtmlToDocx(container, accent_rgb, reuse_para=reuse_para)
        renderer.feed(html or "")
        return renderer.wrote_anything

    def _add_image(container, path, **size_kwargs):
        """container: anything exposing add_picture() directly — a Document
        or a table cell (both add the picture as a new paragraph)."""
        src = _docx_image_source(path)
        if src is None:
            return False
        try:
            container.add_picture(src, **size_kwargs)
            return True
        except Exception:
            log.warning("Could not embed image %s into DOCX", path)
            return False

    def _add_image_to_paragraph(paragraph, path, **size_kwargs):
        """Embeds into an already-existing paragraph (e.g. the banner's first,
        otherwise-empty paragraph) instead of creating a new one."""
        src = _docx_image_source(path)
        if src is None:
            return False
        try:
            paragraph.add_run().add_picture(src, **size_kwargs)
            return True
        except Exception:
            log.warning("Could not embed image %s into DOCX", path)
            return False

    def _parse_page_margin(margin_str):
        parts = (margin_str or "").split()
        try:
            nums = [float(re.sub(r"[a-zA-Z%]+", "", p)) for p in parts if p]
            if len(nums) == 1:
                return nums[0], nums[0]
            if len(nums) >= 2:
                return nums[0], nums[1]
        except Exception:
            pass
        return 18.0, 14.0

    try:
        content = json.loads(quotation_row.CONTENT_JSON)
        info = _resolve_company_info(content, company)
        company_addr = info.get("FULL_ADDRESS") or format_full_address(company)

        letterhead = content.get("letterhead", {})
        style = content.get("style", {})
        accent_hex = style.get("accentColor", "#C8102E")
        accent_rgb = _rgb(accent_hex, "C8102E")
        v_margin, h_margin = _parse_page_margin(style.get("pageMargin", "18mm 14mm"))

        tokens = {
            "customer_name": "",
            "customer_address": "",
            "company_legal_name": info.get("LEGAL_NAME") or "",
            "company_address": company_addr,
            "company_website": info.get("WEBSITE") or "",
        }

        document = docx.Document()

        # Body font — closer to the PDF's Helvetica/Arial than Word's default Calibri.
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10.5)

        section0 = document.sections[0]
        section0.top_margin = Mm(v_margin)
        section0.bottom_margin = Mm(v_margin)
        section0.left_margin = Mm(h_margin)
        section0.right_margin = Mm(h_margin)
        usable_width = section0.page_width - section0.left_margin - section0.right_margin

        # ── Banner header (logo + title + subtitle on an accent-colored band) ──
        banner = document.add_table(rows=1, cols=1)
        banner.autofit = True
        banner_cell = banner.rows[0].cells[0]
        banner_cell.width = usable_width
        _set_cell_shading(banner_cell, accent_hex)

        logo_url = letterhead.get("logoUrlOverride") or info.get("LOGO_URL")
        logo_path = _resolve_static_path(logo_url) if logo_url else None
        first_para = banner_cell.paragraphs[0]
        if logo_path is not None:
            _add_image_to_paragraph(first_para, logo_path, height=Inches(0.45))
            title_para = banner_cell.add_paragraph()
        else:
            title_para = first_para
        title_run = title_para.add_run(letterhead.get("title", "QUOTATION"))
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        if letterhead.get("subtitle"):
            subtitle_para = banner_cell.add_paragraph()
            subtitle_run = subtitle_para.add_run(letterhead["subtitle"])
            subtitle_run.font.size = Pt(9)
            subtitle_run.font.color.rgb = RGBColor(0xF1, 0xF5, 0xF9)

        document.add_paragraph()  # breathing room after the banner

        # ── Meta table: customer info (left) + quote ref/date (right, right-aligned) ──
        meta = document.add_table(rows=1, cols=2)
        meta.autofit = True
        left_cell, right_cell = meta.rows[0].cells
        left_cell.width = usable_width // 2
        right_cell.width = usable_width // 2

        customer = content.get("customerInfo", {})
        if customer.get("showBlock", True):
            lbl_run = left_cell.paragraphs[0].add_run(customer.get("label", "BILL TO"))
            lbl_run.font.size = Pt(8)
            lbl_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            name_run = left_cell.add_paragraph().add_run(customer.get("namePlaceholder", ""))
            name_run.bold = True
            if customer.get("addressPlaceholder"):
                left_cell.add_paragraph(customer["addressPlaceholder"])
            if customer.get("freeText"):
                left_cell.add_paragraph(customer["freeText"])

        right_lines_written = [False]

        def _right_line(text, size=10, color=None, bold=False):
            if not right_lines_written[0]:
                p = right_cell.paragraphs[0]
                right_lines_written[0] = True
            else:
                p = right_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = color

        _right_line("QUOTE REFERENCE NUMBER", size=8, color=RGBColor(0x64, 0x74, 0x8B))
        _right_line(quotation_row.QUOTATION_NUMBER or "", size=11, bold=True)
        _right_line("QUOTE DATE", size=8, color=RGBColor(0x64, 0x74, 0x8B))
        _right_line(str(quotation_row.QUOTATION_DATE or ""), size=10)

        document.add_paragraph()

        # ── Intro ──
        intro_html = _substitute_tokens(content.get("introHtml", ""), tokens)
        _render_html(document, intro_html)

        # ── Sections ──
        for section in sorted(content.get("sections", []), key=lambda s: s.get("order", 0)):
            if section.get("pageBreakBefore"):
                document.add_page_break()
            if section.get("title"):
                title_p = document.add_paragraph()
                title_p.paragraph_format.space_before = Pt(10)
                tr = title_p.add_run(section["title"])
                tr.bold = True
                tr.font.color.rgb = accent_rgb

            stype = section.get("type")
            if stype in ("richtext", "custom"):
                _render_html(document, section.get("html", ""))
            elif stype == "table":
                rows = section.get("rows", [])
                table = document.add_table(rows=1, cols=5)
                table.autofit = True
                hdr = table.rows[0].cells
                headers = ["#", "Product Description", "Qty", "Unit Price", "Amount"]
                for cell, text in zip(hdr, headers):
                    cell.text = ""
                    r = cell.paragraphs[0].add_run(text)
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = _rgb("#8B0B1F")
                    _set_cell_shading(cell, "FEF2F2")
                    _set_cell_border(cell, bottom={"sz": 6, "color": "FECACA"})
                for idx, r in enumerate(rows, start=1):
                    qty = r.get("qty") or 0
                    unit_price = r.get("unitPrice") or 0
                    amount = float(qty) * float(unit_price)
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = r.get("description") or ""
                    row_cells[2].text = str(qty)
                    row_cells[3].text = f"{unit_price:,.2f}"
                    row_cells[4].text = f"{amount:,.2f}"
                    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for c in (row_cells[2], row_cells[3], row_cells[4]):
                        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for c in row_cells:
                        if c.paragraphs[0].runs:
                            c.paragraphs[0].runs[0].font.size = Pt(9.5)
                        _set_cell_border(c, bottom={"sz": 4, "color": "E2E8F0"})
                document.add_paragraph()
            elif stype == "image":
                image_path = _resolve_static_path(section.get("imageUrl"))
                if image_path is not None:
                    _add_image(document, image_path, width=Inches(4))

        # ── Terms & Conditions (shaded box, accent left border) ──
        terms_html = _substitute_tokens(content.get("termsHtml", ""), tokens)
        if terms_html.strip():
            terms_table = document.add_table(rows=1, cols=1)
            terms_table.autofit = True
            terms_cell = terms_table.rows[0].cells[0]
            terms_cell.width = usable_width
            _set_cell_shading(terms_cell, "F8FAFC")
            _set_cell_border(terms_cell, left={"sz": 24, "color": accent_hex.lstrip("#")})
            heading_run = terms_cell.paragraphs[0].add_run("Terms & Conditions")
            heading_run.bold = True
            heading_run.font.color.rgb = accent_rgb
            _render_html(terms_cell, terms_html, accent_rgb=accent_rgb)
            document.add_paragraph()

        # ── Notes ──
        notes_html = _substitute_tokens(content.get("notesHtml", ""), tokens)
        _render_html(document, notes_html)

        # ── Signature ──
        signature = content.get("signature", {})
        if signature.get("showBlock", True) and (signature.get("name") or signature.get("imageUrl")):
            document.add_paragraph("Thanks & Regards,")
            document.add_paragraph(f"For {info.get('LEGAL_NAME') or ''},")
            sig_path = _resolve_static_path(signature.get("imageUrl"))
            if sig_path is not None:
                _add_image(document, sig_path, height=Inches(0.6))
            if signature.get("name"):
                run = document.add_paragraph().add_run(signature["name"])
                run.bold = True
            if signature.get("designation"):
                document.add_paragraph(signature["designation"])

        # ── Footer (small, gray, centered) ──
        footer_html = _substitute_tokens(content.get("footerHtml", ""), tokens)
        footer = document.sections[0].footer
        footer.paragraphs[0].text = ""
        wrote = _render_html(footer, footer_html or f"<p>{company_addr}</p>", reuse_para=footer.paragraphs[0])
        if not wrote:
            footer.paragraphs[0].add_run(company_addr)
        for p in footer.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        buf = BytesIO()
        document.save(buf)
        return buf.getvalue()
    except Exception:
        log.exception("Quotation DOCX render failed")
        return None
